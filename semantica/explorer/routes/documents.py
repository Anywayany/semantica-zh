"""Document-to-graph preview and commit routes for the Explorer."""

from __future__ import annotations

import asyncio
import hashlib
import importlib.util
import io
import logging
import random
import re
import time
import unicodedata
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Dict, Iterable, List, Literal, Optional, Tuple

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from pydantic import BaseModel, Field, SecretStr

from ..dependencies import get_session
from ..session import GraphSession

router = APIRouter(prefix="/api/documents", tags=["Document Graph Builder"])

_DOCUMENT_MAX_BYTES = 25 * 1024 * 1024
_TEXT_MAX_CHARACTERS = 500_000
_CHUNK_SIZE = 4_000
_CHUNK_OVERLAP = 240
_MAX_ENTITIES = 750
_MAX_RELATIONS = 1_500
_LLM_MAX_ATTEMPTS = 3
_DEEPSEEK_TIMEOUT_SECONDS = 120.0
_ALLOWED_EXTENSIONS = frozenset({".txt", ".md", ".markdown", ".html", ".htm", ".docx", ".pdf"})
_SUPPORTED_PROVIDERS = ("openai", "gemini", "groq", "anthropic", "ollama", "deepseek", "novita")
_DEFAULT_MODELS = {
    "openai": "gpt-4o-mini",
    "gemini": "gemini-2.0-flash",
    "groq": "llama-3.3-70b-versatile",
    "anthropic": "claude-3-5-haiku-latest",
    "ollama": "qwen2.5:7b",
    "deepseek": "deepseek-v4-flash",
    "novita": "meta-llama/llama-3.1-8b-instruct",
}

logger = logging.getLogger(__name__)

_ONTOLOGY_PROFILES: Dict[str, Dict[str, Tuple[str, ...]]] = {
    "general": {
        "entity_types": (
            "PERSON", "ORG", "GPE", "LOCATION", "DATE", "EVENT", "PRODUCT",
            "MONEY", "PERCENT", "LAW", "COURT", "DISEASE", "DRUG", "GENE",
        ),
        "relation_types": (
            "works_for", "founded", "located_in", "born_in", "acquired", "part_of",
            "invested_in", "owns", "appointed_as", "founded_on", "develops", "treats",
            "causes", "regulated_by", "decided_by", "signed_on",
        ),
    },
    "business": {
        "entity_types": ("PERSON", "ORG", "GPE", "LOCATION", "DATE", "PRODUCT", "MONEY", "PERCENT"),
        "relation_types": (
            "works_for", "founded", "located_in", "born_in", "acquired", "part_of",
            "invested_in", "owns", "appointed_as", "founded_on", "develops",
        ),
    },
    "biomedical": {
        "entity_types": ("PERSON", "ORG", "GPE", "DATE", "DISEASE", "DRUG", "GENE", "PRODUCT"),
        "relation_types": ("works_for", "located_in", "develops", "treats", "causes", "part_of"),
    },
    "legal": {
        "entity_types": ("PERSON", "ORG", "GPE", "DATE", "LAW", "COURT", "MONEY"),
        "relation_types": ("works_for", "located_in", "part_of", "regulated_by", "decided_by", "signed_on"),
    },
}


class ProviderCapability(BaseModel):
    id: str
    available: bool
    default_model: str
    reason: str
    credential_source: Literal["session", "server", "none", "not_required"] = "none"
    session_configured: bool = False


class DocumentExtractionCapabilities(BaseModel):
    providers: List[ProviderCapability]
    local_nlp_available: bool
    ontology_profiles: Dict[str, Dict[str, List[str]]]


class SessionCredentialRequest(BaseModel):
    provider: str
    api_key: SecretStr


class SessionCredentialResponse(BaseModel):
    provider: str
    configured: bool
    scope: Literal["process_session"] = "process_session"


class LLMExtractionFailure(Exception):
    """Sanitized provider failure with enough context for UI recovery."""

    def __init__(
        self,
        *,
        code: str,
        provider: str,
        model: str,
        phase: str,
        chunk_index: int,
        total_chunks: int,
        attempts: int,
        retryable: bool,
        upstream_status: Optional[int] = None,
    ) -> None:
        super().__init__(code)
        self.code = code
        self.provider = provider
        self.model = model
        self.phase = phase
        self.chunk_index = chunk_index
        self.total_chunks = total_chunks
        self.attempts = attempts
        self.retryable = retryable
        self.upstream_status = upstream_status

    def detail(self) -> Dict[str, Any]:
        return {
            "code": self.code,
            "provider": self.provider,
            "model": self.model,
            "phase": self.phase,
            "chunk": self.chunk_index,
            "total_chunks": self.total_chunks,
            "attempts": self.attempts,
            "retryable": self.retryable,
            "upstream_status": self.upstream_status,
        }


class ExtractionExecution(BaseModel):
    requested_method: Literal["auto", "rules", "llm"]
    actual_methods: List[str] = Field(default_factory=list)
    provider: Optional[str] = None
    model: Optional[str] = None
    provider_available: bool = False
    fallback_used: bool = False
    ontology_profile: str = "general"


class DocumentEntityCandidate(BaseModel):
    id: str
    text: str
    label: str
    confidence: float = Field(ge=0.0, le=1.0)
    mention_count: int = Field(default=1, ge=1)
    properties: Dict[str, Any] = Field(default_factory=dict)


class DocumentRelationCandidate(BaseModel):
    id: str
    source_id: str
    target_id: str
    predicate: str
    confidence: float = Field(ge=0.0, le=1.0)
    context: str = ""
    mention_count: int = Field(default=1, ge=1)
    properties: Dict[str, Any] = Field(default_factory=dict)


class DocumentPreviewResponse(BaseModel):
    document_id: str
    filename: str
    media_type: str
    parser: str
    language: Literal["en", "zh"]
    extraction_method: str
    execution: ExtractionExecution
    ontology_profile: str = "general"
    character_count: int
    chunk_count: int
    text_preview: str
    entities: List[DocumentEntityCandidate] = Field(default_factory=list)
    relations: List[DocumentRelationCandidate] = Field(default_factory=list)
    warnings: List[str] = Field(default_factory=list)


class DocumentCommitRequest(BaseModel):
    document_id: str
    filename: str
    media_type: str = "application/octet-stream"
    parser: str
    language: Literal["en", "zh"]
    extraction_method: str
    ontology_profile: str = "general"
    character_count: int = Field(default=0, ge=0)
    entities: List[DocumentEntityCandidate] = Field(default_factory=list, max_length=_MAX_ENTITIES)
    relations: List[DocumentRelationCandidate] = Field(default_factory=list, max_length=_MAX_RELATIONS)


class DocumentCommitResponse(BaseModel):
    status: Literal["success"] = "success"
    document_node_id: str
    nodes_added: int
    edges_added: int
    entities_submitted: int
    relations_submitted: int


class _VisibleHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: List[str] = []
        self._hidden_depth = 0

    def handle_starttag(self, tag: str, _attrs: List[Tuple[str, Optional[str]]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._hidden_depth += 1
        elif tag in {"p", "br", "div", "section", "article", "li", "h1", "h2", "h3", "h4", "tr"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._hidden_depth:
            self._hidden_depth -= 1
        elif tag in {"p", "div", "section", "article", "li", "h1", "h2", "h3", "h4", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._hidden_depth:
            self.parts.append(data)


def _provider_runtime_status(
    provider: str,
    model: Optional[str] = None,
    session_api_key: Optional[str] = None,
) -> ProviderCapability:
    provider = provider.strip().lower()
    if provider not in _SUPPORTED_PROVIDERS:
        return ProviderCapability(
            id=provider,
            available=False,
            default_model="",
            reason="Unsupported provider.",
        )
    resolved_model = (model or "").strip() or _DEFAULT_MODELS[provider]
    server_api_key: Optional[str] = None
    credential_source: Literal["session", "server", "none", "not_required"] = "not_required"
    if provider != "ollama":
        from ...semantic_extract.config import config

        server_api_key = config.get_api_key(provider)
        credential_source = "session" if session_api_key else "server" if server_api_key else "none"
    dependency = {
        "openai": "openai",
        "gemini": "google",
        "groq": "groq",
        "anthropic": "anthropic",
        "ollama": "ollama",
        "deepseek": "openai",
        "novita": "openai",
    }[provider]
    if importlib.util.find_spec(dependency) is None:
        return ProviderCapability(
            id=provider,
            available=False,
            default_model=_DEFAULT_MODELS[provider],
            reason=f"Python dependency '{dependency}' is not installed.",
            credential_source=credential_source,
            session_configured=bool(session_api_key),
        )
    try:
        if provider == "ollama":
            from ...semantic_extract.providers import create_provider

            runtime = create_provider(provider, use_pool=False, model=resolved_model)
            available = bool(runtime.is_available())
            reason = "Ready" if available else "Ollama is not reachable."
        else:
            available = bool(session_api_key or server_api_key)
            reason = (
                "Configured for this server session; connectivity is verified during extraction."
                if session_api_key
                else "Configured; connectivity is verified during extraction."
                if server_api_key
                else "Server-side API key is missing."
            )
    except Exception as exc:
        return ProviderCapability(
            id=provider,
            available=False,
            default_model=_DEFAULT_MODELS[provider],
            reason=f"Provider initialization failed: {type(exc).__name__}",
            credential_source=credential_source,
            session_configured=bool(session_api_key),
        )
    return ProviderCapability(
        id=provider,
        available=available,
        default_model=_DEFAULT_MODELS[provider],
        reason=reason,
        credential_source=credential_source,
        session_configured=bool(session_api_key),
    )


def _local_nlp_available() -> bool:
    return importlib.util.find_spec("spacy") is not None and importlib.util.find_spec("en_core_web_sm") is not None


def _parse_type_list(value: Optional[str]) -> Tuple[str, ...]:
    if not value:
        return ()
    normalized = []
    for item in re.split(r"[,，;；\n]+", value):
        candidate = re.sub(r"[^A-Za-z0-9_:-]", "", item.strip()).upper()
        if candidate and candidate not in normalized:
            normalized.append(candidate)
    return tuple(normalized[:100])


def _resolve_ontology(
    profile: str,
    custom_entity_types: Optional[str],
    custom_relation_types: Optional[str],
) -> Tuple[str, Tuple[str, ...], Tuple[str, ...]]:
    profile = profile.strip().lower()
    if profile == "custom":
        entity_types = _parse_type_list(custom_entity_types)
        relation_types = tuple(item.lower() for item in _parse_type_list(custom_relation_types))
        if not entity_types or not relation_types:
            raise HTTPException(
                status_code=422,
                detail="Custom ontology requires at least one entity type and one relation type.",
            )
        return profile, entity_types, relation_types
    definition = _ONTOLOGY_PROFILES.get(profile)
    if definition is None:
        raise HTTPException(
            status_code=422,
            detail=f"Unknown ontology profile '{profile}'.",
        )
    return profile, definition["entity_types"], definition["relation_types"]


def _session_credentials(request: Request) -> Dict[str, str]:
    credentials = getattr(request.app.state, "llm_session_credentials", None)
    if credentials is None:
        credentials = {}
        request.app.state.llm_session_credentials = credentials
    return credentials


def _require_local_credential_management(request: Request) -> None:
    client_host = request.client.host if request.client else ""
    if client_host not in {"127.0.0.1", "::1", "localhost", "testclient"}:
        raise HTTPException(
            status_code=403,
            detail="Session credentials can only be managed from the local machine.",
        )


@router.post("/providers/session", response_model=SessionCredentialResponse)
async def configure_session_credential(
    payload: SessionCredentialRequest,
    request: Request,
) -> SessionCredentialResponse:
    _require_local_credential_management(request)
    provider = payload.provider.strip().lower()
    if provider not in _SUPPORTED_PROVIDERS or provider == "ollama":
        raise HTTPException(status_code=422, detail=f"Unsupported credential provider '{provider}'.")
    api_key = payload.api_key.get_secret_value().strip()
    if not api_key:
        raise HTTPException(status_code=422, detail="API key must not be empty.")
    if len(api_key) > 8_192:
        raise HTTPException(status_code=422, detail="API key is too long.")
    _session_credentials(request)[provider] = api_key
    return SessionCredentialResponse(provider=provider, configured=True)


@router.delete("/providers/session/{provider}", response_model=SessionCredentialResponse)
async def clear_session_credential(provider: str, request: Request) -> SessionCredentialResponse:
    _require_local_credential_management(request)
    provider = provider.strip().lower()
    if provider not in _SUPPORTED_PROVIDERS or provider == "ollama":
        raise HTTPException(status_code=422, detail=f"Unsupported credential provider '{provider}'.")
    _session_credentials(request).pop(provider, None)
    return SessionCredentialResponse(provider=provider, configured=False)


@router.get("/capabilities", response_model=DocumentExtractionCapabilities)
async def document_extraction_capabilities(request: Request) -> DocumentExtractionCapabilities:
    session_credentials = dict(_session_credentials(request))
    providers = await asyncio.gather(*[
        asyncio.to_thread(_provider_runtime_status, provider, None, session_credentials.get(provider))
        for provider in _SUPPORTED_PROVIDERS
    ])
    ontology_profiles = {
        name: {
            "entity_types": list(definition["entity_types"]),
            "relation_types": list(definition["relation_types"]),
        }
        for name, definition in _ONTOLOGY_PROFILES.items()
    }
    ontology_profiles["custom"] = {"entity_types": [], "relation_types": []}
    return DocumentExtractionCapabilities(
        providers=list(providers),
        local_nlp_available=_local_nlp_available(),
        ontology_profiles=ontology_profiles,
    )


def _safe_filename(filename: Optional[str]) -> str:
    normalized = (filename or "document").replace("\\", "/")
    return normalized.rsplit("/", 1)[-1] or "document"


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=422, detail="Document text must use UTF-8, UTF-16, or GB18030 encoding.")


def _normalize_document_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_document(filename: str, content: bytes) -> Tuple[str, str, List[str]]:
    extension = Path(filename).suffix.lower()
    warnings: List[str] = []

    if extension in {".txt", ".md", ".markdown"}:
        return _normalize_document_text(_decode_text(content)), "plain-text", warnings

    if extension in {".html", ".htm"}:
        parser = _VisibleHTMLParser()
        parser.feed(_decode_text(content))
        return _normalize_document_text("".join(parser.parts)), "html-parser", warnings

    if extension == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise HTTPException(
                status_code=503,
                detail="DOCX parsing requires python-docx. Install semantica[explorer] dependencies.",
            ) from exc

        document = Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
        return _normalize_document_text("\n\n".join(parts)), "python-docx", warnings

    if extension == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise HTTPException(
                status_code=503,
                detail="PDF parsing requires pypdf. Install semantica[explorer] dependencies.",
            ) from exc

        try:
            reader = PdfReader(io.BytesIO(content))
            pages = [(page.extract_text() or "").strip() for page in reader.pages]
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Unable to parse PDF: {exc}") from exc
        if pages and not any(pages):
            warnings.append("The PDF contains no extractable text; scanned documents require OCR/Docling.")
        return _normalize_document_text("\n\n".join(page for page in pages if page)), "pypdf", warnings

    raise HTTPException(
        status_code=422,
        detail=f"Unsupported document type '{extension}'. Allowed: {sorted(_ALLOWED_EXTENSIONS)}",
    )


def _detect_language(text: str, requested: str) -> Literal["en", "zh"]:
    if requested in {"en", "zh"}:
        return requested
    sample = text[:20_000]
    cjk_count = len(re.findall(r"[\u3400-\u9fff]", sample))
    letter_count = len(re.findall(r"[A-Za-z\u3400-\u9fff]", sample)) or 1
    return "zh" if cjk_count / letter_count >= 0.2 else "en"


def _chunk_text(text: str) -> List[Tuple[int, str]]:
    if len(text) <= _CHUNK_SIZE:
        return [(0, text)] if text else []

    chunks: List[Tuple[int, str]] = []
    start = 0
    while start < len(text):
        hard_end = min(len(text), start + _CHUNK_SIZE)
        end = hard_end
        if hard_end < len(text):
            boundary_floor = start + int(_CHUNK_SIZE * 0.65)
            candidates = [
                text.rfind(token, boundary_floor, hard_end)
                for token in ("\n\n", "。", "！", "？", ". ", "! ", "? ")
            ]
            boundary = max(candidates)
            if boundary >= boundary_floor:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append((start, chunk))
        if end >= len(text):
            break
        next_start = max(start + 1, end - _CHUNK_OVERLAP)
        start = next_start
    return chunks


def _normalize_entity_key(label: str, text: str) -> Tuple[str, str]:
    normalized = unicodedata.normalize("NFKC", text).strip().casefold()
    return label.strip().upper() or "ENTITY", normalized


def _entity_id(label: str, text: str) -> str:
    key = "\0".join(_normalize_entity_key(label, text)).encode("utf-8")
    return f"entity:{hashlib.sha1(key).hexdigest()[:16]}"


def _relation_id(source_id: str, predicate: str, target_id: str) -> str:
    key = f"{source_id}\0{predicate.strip().casefold()}\0{target_id}".encode("utf-8")
    return f"relation:{hashlib.sha1(key).hexdigest()[:16]}"


def _entity_canonical_text(entity: Any) -> str:
    metadata = getattr(entity, "metadata", {}) or {}
    return str(metadata.get("canonical_text") or getattr(entity, "text", "")).strip()


def _cjk_entities(text: str, allowed_entity_types: Tuple[str, ...]) -> List[Any]:
    from ...semantic_extract.types import Entity

    entities: List[Entity] = []
    covered: List[Tuple[int, int]] = []
    allowed = set(allowed_entity_types)
    aliases_by_canonical: Dict[str, set[str]] = {}

    def add_candidate(
        label: str,
        start: int,
        end: int,
        confidence: float,
        rule: str,
        canonical_text: Optional[str] = None,
    ) -> None:
        if label not in allowed:
            return
        candidate = text[start:end].strip()
        leading_space = len(text[start:end]) - len(text[start:end].lstrip())
        start += leading_space
        end = start + len(candidate)
        if len(candidate) < 2:
            return
        if label == "ORG" and candidate in {"该公司", "本公司", "其公司", "公司", "该机构", "本机构"}:
            return
        if any(existing_start < end and existing_end > start for existing_start, existing_end in covered):
            return
        canonical = (canonical_text or candidate).strip()
        aliases = sorted(aliases_by_canonical.get(canonical, set()))
        entities.append(
            Entity(
                text=candidate,
                label=label,
                start_char=start,
                end_char=end,
                confidence=confidence,
                metadata={
                    "extraction_method": "cjk-rules-v2",
                    "canonical_text": canonical,
                    "aliases": aliases,
                    "rule": rule,
                    "confidence_factors": {
                        "rule": rule,
                        "raw_score": confidence,
                        "surface_match": True,
                    },
                },
            )
        )
        covered.append((start, end))

    delimiters = "，。！？；：、\n\t ()（）[]【】"
    boundary_tokens = (
        "创立", "创建", "成立", "加入", "任职于", "就职于", "担任", "位于", "坐落于",
        "出生于", "收购", "隶属于", "投资", "控股", "研发", "开发", "治疗", "来自", "在",
        "由", "被", "与", "和",
    )

    def suffix_mentions(label: str, suffix_pattern: str, max_prefix: int, confidence: float) -> None:
        for match in re.finditer(suffix_pattern, text):
            window_start = max(0, match.start() - max_prefix)
            start = window_start
            prefix = text[window_start:match.start()]
            for delimiter in delimiters:
                index = prefix.rfind(delimiter)
                if index >= 0:
                    start = max(start, window_start + index + 1)
            for token in boundary_tokens:
                index = prefix.rfind(token)
                if index >= 0:
                    start = max(start, window_start + index + len(token))
            add_candidate(label, start, match.end(), confidence, f"{label.lower()}_suffix")

    suffix_mentions(
        "ORG",
        r"股份有限公司|有限责任公司|有限公司|集团|公司|大学|学院|研究院|研究所|委员会|银行|医院",
        28,
        0.88,
    )
    suffix_mentions("GPE", r"自治区|特别行政区|省|市|区|县|国", 12, 0.86)

    for entity in list(entities):
        if entity.label != "ORG":
            continue
        suffix = text[entity.end_char:entity.end_char + 45]
        alias_match = re.match(
            r"\s*[（(](?:以下简称|简称)\s*[“\"']?(?P<alias>[\u3400-\u9fffA-Za-z0-9·_-]{2,24})",
            suffix,
        )
        if alias_match:
            canonical = _entity_canonical_text(entity)
            alias = alias_match.group("alias")
            aliases_by_canonical.setdefault(canonical, set()).add(alias)
            entity.metadata["aliases"] = sorted(aliases_by_canonical[canonical])

    for canonical, aliases in aliases_by_canonical.items():
        for alias in aliases:
            for match in re.finditer(re.escape(alias), text):
                add_candidate("ORG", match.start(), match.end(), 0.82, "declared_alias", canonical)

    known_places = (
        "北京", "上海", "天津", "重庆", "香港", "澳门", "台湾", "中国", "美国", "英国",
        "法国", "德国", "日本", "韩国", "新加坡", "深圳", "广州", "杭州", "南京", "武汉",
        "成都", "西安", "苏州",
    )
    for place in known_places:
        for match in re.finditer(re.escape(place), text):
            add_candidate("GPE", match.start(), match.end(), 0.84, "place_lexicon")

    for label, pattern in (
        ("DATE", r"(?:\d{4}年(?:\d{1,2}月(?:\d{1,2}日)?)?|\d{1,2}月\d{1,2}日)"),
        ("MONEY", r"(?:人民币)?[￥¥]?\d+(?:\.\d+)?(?:万|亿)?元"),
        ("PERCENT", r"\d+(?:\.\d+)?%"),
    ):
        for match in re.finditer(pattern, text):
            add_candidate(label, match.start(), match.end(), 0.95, f"{label.lower()}_format")

    person_patterns = (
        r"(?P<person>[\u3400-\u9fff·]{2,6})(?=\s*(?:创立|创建|加入|任职于|就职于|担任|出任|出生于|投资|收购))",
        r"(?:董事长|总经理|首席执行官|CEO|教授|医生|律师)\s*(?P<person>[\u3400-\u9fff·]{2,4})",
        r"(?P<person>[\u3400-\u9fff·]{2,4})(?=\s*(?:先生|女士|博士|教授|律师))",
        r"(?:由|被)\s*(?P<person>[\u3400-\u9fff·]{2,4})(?=\s*(?:创立|创建|收购|担任|负责))",
    )
    person_stopwords = {"该公司", "本公司", "公司", "集团", "其公司", "该机构", "本机构"}
    for pattern in person_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            person = match.group("person")
            start, end = match.span("person")
            while person and person[0] in "由被与和及":
                person = person[1:]
                start += 1
            if person in person_stopwords or person.endswith(("公司", "集团", "大学", "银行", "医院")):
                continue
            add_candidate("PERSON", start, end, 0.84, "person_context")

    domain_patterns = (
        ("LAW", r"《[^》\n]{2,60}》", 0.94, "law_title"),
        ("COURT", r"[^，。！？；：\n]{2,30}(?:人民法院|法院)", 0.88, "court_suffix"),
        ("DISEASE", r"[^，。！？；：\n]{2,18}(?:综合征|癌症|癌|疾病|感染|炎)", 0.78, "disease_suffix"),
        ("DRUG", r"[^，。！？；：\n]{2,16}(?:注射液|胶囊|片剂|药物|药)", 0.76, "drug_suffix"),
    )
    for label, pattern, confidence, rule in domain_patterns:
        for match in re.finditer(pattern, text):
            add_candidate(label, match.start(), match.end(), confidence, rule)
    return sorted(entities, key=lambda entity: entity.start_char)


def _cjk_relations(
    text: str,
    entities: List[Any],
    allowed_relation_types: Tuple[str, ...],
) -> List[Any]:
    from ...semantic_extract.types import Relation

    predicate_tokens = {
        "任职于": "works_for",
        "就职于": "works_for",
        "加入": "works_for",
        "担任": "works_for",
        "出任": "works_for",
        "创立": "founded",
        "创建": "founded",
        "成立于": "founded_on",
        "位于": "located_in",
        "坐落于": "located_in",
        "出生于": "born_in",
        "收购": "acquired",
        "隶属于": "part_of",
        "投资": "invested_in",
        "控股": "owns",
        "研发": "develops",
        "开发": "develops",
        "治疗": "treats",
        "导致": "causes",
        "受监管于": "regulated_by",
        "签署于": "signed_on",
    }
    allowed = set(allowed_relation_types)
    relations: List[Relation] = []
    sentence_ranges = [match.span() for match in re.finditer(r"[^。！？!?\n]+[。！？!?\n]?", text)]
    prior_entities: List[Any] = []

    def last_prior(*labels: str) -> Optional[Any]:
        expected = set(labels)
        for entity in reversed(prior_entities):
            if entity.label in expected:
                return entity
        return None

    def add_relation(
        subject: Any,
        predicate: str,
        obj: Any,
        confidence: float,
        sentence: str,
        token: str,
        resolution: str,
    ) -> None:
        if predicate not in allowed or _entity_canonical_text(subject) == _entity_canonical_text(obj):
            return
        relations.append(
            Relation(
                subject=subject,
                predicate=predicate,
                object=obj,
                confidence=confidence,
                context=sentence.strip(),
                metadata={
                    "extraction_method": "cjk-rules-v2",
                    "trigger": token,
                    "evidence": sentence.strip(),
                    "resolution": resolution,
                    "confidence_factors": {
                        "explicit_trigger": True,
                        "surface_evidence": True,
                        "coreference": resolution == "pronoun",
                        "raw_score": confidence,
                    },
                },
            )
        )

    for sentence_start, sentence_end in sentence_ranges:
        sentence_entities = [
            entity for entity in entities
            if entity.start_char >= sentence_start and entity.end_char <= sentence_end
        ]
        sentence = text[sentence_start:sentence_end]
        for token, predicate in sorted(predicate_tokens.items(), key=lambda item: -len(item[0])):
            for token_match in re.finditer(re.escape(token), sentence):
                token_index = token_match.start()
                global_index = sentence_start + token_index
                before = [entity for entity in sentence_entities if entity.end_char <= global_index]
                after = [entity for entity in sentence_entities if entity.start_char >= global_index + len(token)]

                marker_matches = list(re.finditer(r"由|被", sentence[:token_index]))
                if predicate in {"founded", "acquired"} and marker_matches:
                    marker_index = marker_matches[-1].start()
                    marker_global = sentence_start + marker_index
                    agents = [entity for entity in before if entity.start_char >= marker_global + 1]
                    patients = [entity for entity in before if entity.end_char <= marker_global]
                    if agents and patients:
                        add_relation(
                            max(agents, key=lambda entity: entity.end_char),
                            predicate,
                            max(patients, key=lambda entity: entity.end_char),
                            0.91,
                            sentence,
                            token,
                            "passive",
                        )
                        continue

                subject = max(before, key=lambda entity: entity.end_char) if before else None
                obj = min(after, key=lambda entity: entity.start_char) if after else None
                resolution = "surface"
                prefix = sentence[:token_index].rstrip()
                if subject is None and re.search(r"(?:该公司|本公司|该机构|本机构|其总部|公司)$", prefix):
                    subject = last_prior("ORG")
                    resolution = "pronoun"
                elif subject is None and re.search(r"(?:他|她|该人士|其本人)$", prefix):
                    subject = last_prior("PERSON")
                    resolution = "pronoun"
                if subject is None or obj is None:
                    continue
                add_relation(
                    subject,
                    predicate,
                    obj,
                    0.74 if resolution == "pronoun" else 0.88,
                    sentence,
                    token,
                    resolution,
                )
        prior_entities.extend(sentence_entities)
    return relations


def _exception_status(exc: BaseException) -> Optional[int]:
    current: Optional[BaseException] = exc
    seen: set[int] = set()
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        status = getattr(current, "status_code", None)
        if isinstance(status, int):
            return status
        response = getattr(current, "response", None)
        response_status = getattr(response, "status_code", None)
        if isinstance(response_status, int):
            return response_status
        current = current.__cause__ or current.__context__
    return None


def _classify_llm_error(exc: BaseException) -> Tuple[str, bool, Optional[int]]:
    status = _exception_status(exc)
    message = str(exc).lower()
    if status == 429 or "rate limit" in message or "too many requests" in message:
        return "provider_rate_limited", True, status or 429
    if status in {408, 409} or "timeout" in message or "timed out" in message:
        return "provider_timeout", True, status
    if status is not None and status >= 500:
        return "provider_service_error", True, status
    if "connection error" in message or "connection reset" in message or "connection refused" in message:
        return "provider_connection_error", True, status
    if "empty response" in message or "empty structured response" in message:
        return "provider_empty_response", True, status
    if status in {401, 403} or "authentication" in message or "invalid api key" in message:
        return "provider_authentication_error", False, status
    if status in {400, 404, 422} or "model not found" in message or "does not exist" in message:
        return "provider_request_error", False, status
    return "provider_structured_output_error", False, status


def _sleep_before_llm_retry(attempt: int) -> None:
    delay = (2 ** (attempt - 1)) + random.uniform(0.0, 0.25)
    time.sleep(delay)


def _create_request_provider(provider: str, model: str, api_key: Optional[str]) -> Any:
    from ...semantic_extract.providers import create_provider

    provider_kwargs: Dict[str, Any] = {"model": model}
    if api_key:
        provider_kwargs["api_key"] = api_key
    if provider == "deepseek":
        # Application-level chunk retries own retry behavior. Keeping SDK retries
        # disabled prevents nested retries and duplicate billable requests.
        provider_kwargs.update(
            timeout=_DEEPSEEK_TIMEOUT_SECONDS,
            max_retries=0,
        )
    return create_provider(provider, use_pool=False, **provider_kwargs)


def _close_request_provider(provider_instance: Optional[Any]) -> None:
    client = getattr(provider_instance, "client", None)
    close = getattr(client, "close", None)
    if callable(close):
        close()


def _llm_call_options(provider: str, model: str, provider_instance: Any) -> Dict[str, Any]:
    options: Dict[str, Any] = {
        "provider_instance": provider_instance,
        "max_retries": 1,
        "fallback_to_manual": False,
    }
    if provider == "deepseek" and model.startswith("deepseek-v4"):
        options.update(
            response_format={"type": "json_object"},
            extra_body={"thinking": {"type": "disabled"}},
        )
    return options


def _extract_chunk(
    text: str,
    language: Literal["en", "zh"],
    extraction_method: str,
    min_confidence: float,
    provider: str,
    model: str,
    provider_available: bool,
    provider_instance: Optional[Any],
    entity_types: Tuple[str, ...],
    relation_types: Tuple[str, ...],
    chunk_index: int = 1,
    total_chunks: int = 1,
) -> Tuple[List[Any], List[Any], set[str], List[str]]:
    from ...semantic_extract.methods import (
        extract_entities_llm,
        extract_entities_pattern,
        extract_relations_llm,
        extract_relations_pattern,
    )

    if extraction_method not in {"auto", "rules", "llm"}:
        raise HTTPException(status_code=422, detail="extraction_method must be auto, rules, or llm")

    warnings: List[str] = []

    def strict_llm() -> Tuple[List[Any], List[Any]]:
        if provider_instance is None:
            raise LLMExtractionFailure(
                code="provider_initialization_error",
                provider=provider,
                model=model,
                phase="provider_initialization",
                chunk_index=chunk_index,
                total_chunks=total_chunks,
                attempts=1,
                retryable=False,
            )
        call_options = _llm_call_options(provider, model, provider_instance)
        try:
            entities = list(extract_entities_llm(
                text,
                provider=provider,
                model=model,
                language=language,
                entity_types=list(entity_types),
                structured_output_mode="typed",
                silent_fail=False,
                **call_options,
            ) or [])
        except Exception as exc:
            code, retryable, upstream_status = _classify_llm_error(exc)
            raise LLMExtractionFailure(
                code=code,
                provider=provider,
                model=model,
                phase="entity_extraction",
                chunk_index=chunk_index,
                total_chunks=total_chunks,
                attempts=1,
                retryable=retryable,
                upstream_status=upstream_status,
            ) from exc
        try:
            relations = list(extract_relations_llm(
                text,
                entities,
                provider=provider,
                model=model,
                language=language,
                relation_types=list(relation_types),
                structured_output_mode="typed",
                silent_fail=False,
                **call_options,
            ) or []) if entities else []
            return entities, relations
        except Exception as exc:
            code, retryable, upstream_status = _classify_llm_error(exc)
            raise LLMExtractionFailure(
                code=code,
                provider=provider,
                model=model,
                phase="relation_extraction",
                chunk_index=chunk_index,
                total_chunks=total_chunks,
                attempts=1,
                retryable=retryable,
                upstream_status=upstream_status,
            ) from exc

    def llm_with_retry() -> Tuple[List[Any], List[Any]]:
        last_failure: Optional[LLMExtractionFailure] = None
        for attempt in range(1, _LLM_MAX_ATTEMPTS + 1):
            try:
                return strict_llm()
            except LLMExtractionFailure as failure:
                failure.attempts = attempt
                last_failure = failure
                logger.warning(
                    "LLM extraction attempt failed provider=%s model=%s phase=%s chunk=%s/%s attempt=%s code=%s retryable=%s",
                    provider,
                    model,
                    failure.phase,
                    chunk_index,
                    total_chunks,
                    attempt,
                    failure.code,
                    failure.retryable,
                )
                if not failure.retryable or attempt >= _LLM_MAX_ATTEMPTS:
                    raise
                _sleep_before_llm_retry(attempt)
        assert last_failure is not None
        raise last_failure

    if extraction_method == "llm":
        if not provider_available:
            raise HTTPException(
                status_code=503,
                detail=f"LLM provider '{provider}' is unavailable. Configure credentials/dependencies or select Automatic/Offline rules.",
            )
        entities, relations = llm_with_retry()
        return entities, relations, {"llm-structured"}, warnings

    if extraction_method == "auto" and provider_available:
        try:
            entities, relations = llm_with_retry()
            if entities:
                return entities, relations, {"llm-structured"}, warnings
            warnings.append("LLM returned no entities; Automatic mode used the offline extractor.")
        except LLMExtractionFailure as exc:
            warnings.append(
                f"LLM failed with {exc.code} in chunk {chunk_index}/{total_chunks}; "
                "Automatic mode used the offline extractor."
            )

    if language == "zh":
        entities = _cjk_entities(text, entity_types)
        relations = _cjk_relations(text, entities, relation_types)
        return entities, relations, {"cjk-rules-v2"}, warnings

    if extraction_method == "auto" and _local_nlp_available():
        from ...semantic_extract.ner_extractor import NERExtractor

        ner = NERExtractor(method="ml", min_confidence=0.0, language=language)
        entities = list(ner.extract_entities(text) or [])
        actual_method = "spacy"
    else:
        entities = list(extract_entities_pattern(text) or [])
        actual_method = "pattern"
    relations = list(extract_relations_pattern(text, entities) or [])
    return entities, relations, {actual_method, "pattern-relations"}, warnings


def _candidate_method(metadata: Dict[str, Any], fallback: str) -> str:
    method = str(metadata.get("extraction_method") or fallback)
    return "llm-structured" if method.startswith("llm") else method


def _calibrate_confidence(raw_score: float, method: str, evidence_strength: float) -> Tuple[float, Dict[str, Any]]:
    reliability = {
        "llm-structured": 0.88,
        "cjk-rules-v2": 0.84,
        "ml": 0.86,
        "spacy": 0.86,
        "pattern": 0.68,
        "pattern-relations": 0.68,
    }.get(method, 0.65)
    raw = max(0.0, min(1.0, raw_score))
    evidence = max(0.0, min(1.0, evidence_strength))
    calibrated = max(0.0, min(1.0, 0.55 * raw + 0.35 * reliability + 0.10 * evidence))
    return round(calibrated, 4), {
        "raw_score": round(raw, 4),
        "method_reliability": reliability,
        "evidence_strength": evidence,
        "formula": "0.55*raw + 0.35*method + 0.10*evidence",
    }


def _evidence_sentence(text: str, subject: str, obj: str, supplied: str) -> str:
    if supplied and subject in supplied and obj in supplied:
        return supplied.strip()
    for match in re.finditer(r"[^。！？!?\n]+[。！？!?\n]?", text):
        sentence = match.group(0).strip()
        if subject in sentence and obj in sentence:
            return sentence
    return ""


def _merge_candidates(
    chunks: Iterable[Tuple[int, str]],
    language: Literal["en", "zh"],
    extraction_method: str,
    min_confidence: float,
    provider: str,
    model: str,
    provider_available: bool,
    provider_instance: Optional[Any],
    entity_types: Tuple[str, ...],
    relation_types: Tuple[str, ...],
) -> Tuple[List[DocumentEntityCandidate], List[DocumentRelationCandidate], List[str], List[str]]:
    entity_records: Dict[Tuple[str, str], Dict[str, Any]] = {}
    relation_records: Dict[Tuple[str, str, str], Dict[str, Any]] = {}
    warnings: List[str] = []
    methods_used: set[str] = set()
    chunk_list = list(chunks)
    total_chunks = len(chunk_list)

    for chunk_offset, (chunk_start, chunk_text) in enumerate(chunk_list):
        entities, relations, chunk_methods, chunk_warnings = _extract_chunk(
            chunk_text,
            language,
            extraction_method,
            min_confidence,
            provider,
            model,
            provider_available,
            provider_instance,
            entity_types,
            relation_types,
            chunk_index=chunk_offset + 1,
            total_chunks=total_chunks,
        )
        methods_used.update(chunk_methods)
        warnings.extend(chunk_warnings)

        for entity in entities:
            label = str(getattr(entity, "label", "ENTITY") or "ENTITY")
            surface_text = str(getattr(entity, "text", "")).strip()
            metadata = getattr(entity, "metadata", {}) or {}
            entity_text = str(metadata.get("canonical_text") or surface_text).strip()
            if not entity_text or label.upper() not in entity_types:
                continue
            method = _candidate_method(metadata, next(iter(chunk_methods), extraction_method))
            confidence, confidence_breakdown = _calibrate_confidence(
                float(getattr(entity, "confidence", 1.0)),
                method,
                1.0 if surface_text and surface_text in chunk_text else 0.0,
            )
            if confidence < min_confidence:
                continue
            key = _normalize_entity_key(label, entity_text)
            record = entity_records.setdefault(
                key,
                {
                    "id": _entity_id(label, entity_text),
                    "text": entity_text,
                    "label": label.upper(),
                    "confidence": confidence,
                    "mention_count": 0,
                    "offsets": [],
                    "evidence": [],
                    "aliases": set(),
                    "methods": set(),
                    "confidence_breakdown": confidence_breakdown,
                },
            )
            record["mention_count"] += 1
            if confidence > record["confidence"]:
                record["confidence"] = confidence
                record["confidence_breakdown"] = confidence_breakdown
            if surface_text and surface_text != entity_text:
                record["aliases"].add(surface_text)
            record["aliases"].update(str(alias) for alias in metadata.get("aliases", []) if alias)
            local_start = int(getattr(entity, "start_char", 0) or 0)
            local_end = int(getattr(entity, "end_char", local_start + len(entity_text)) or local_start + len(entity_text))
            if len(record["offsets"]) < 20:
                record["offsets"].append({
                    "start": chunk_start + local_start,
                    "end": chunk_start + local_end,
                        "chunk": chunk_offset,
                })
            if len(record["evidence"]) < 5:
                evidence_start = max(0, local_start - 60)
                evidence_end = min(len(chunk_text), local_end + 60)
                record["evidence"].append(chunk_text[evidence_start:evidence_end].strip())
            record["methods"].add(method)

        for relation in relations:
            subject = getattr(relation, "subject", None)
            obj = getattr(relation, "object", None)
            if subject is None or obj is None:
                continue
            subject_key = _normalize_entity_key(str(subject.label), _entity_canonical_text(subject))
            object_key = _normalize_entity_key(str(obj.label), _entity_canonical_text(obj))
            if subject_key not in entity_records or object_key not in entity_records:
                continue
            source_id = entity_records[subject_key]["id"]
            target_id = entity_records[object_key]["id"]
            if source_id == target_id:
                continue
            predicate = str(getattr(relation, "predicate", "related_to") or "related_to").strip()
            if predicate.lower() not in relation_types or predicate.lower() == "related_to":
                continue
            key = (source_id, predicate.casefold(), target_id)
            metadata = getattr(relation, "metadata", {}) or {}
            method = _candidate_method(metadata, next(iter(chunk_methods), extraction_method))
            context = _evidence_sentence(
                chunk_text,
                str(getattr(subject, "text", "")),
                str(getattr(obj, "text", "")),
                str(getattr(relation, "context", "") or ""),
            )
            coreference_evidence = metadata.get("resolution") == "pronoun" and bool(metadata.get("evidence"))
            if not context and not coreference_evidence:
                warnings.append(f"Dropped unsupported relation '{predicate}' because no source evidence was found.")
                continue
            if not context:
                context = str(metadata.get("evidence", "")).strip()
            confidence, confidence_breakdown = _calibrate_confidence(
                float(getattr(relation, "confidence", 1.0)),
                method,
                0.8 if coreference_evidence else 1.0,
            )
            if confidence < min_confidence:
                continue
            record = relation_records.setdefault(
                key,
                {
                    "id": _relation_id(source_id, predicate, target_id),
                    "source_id": source_id,
                    "target_id": target_id,
                    "predicate": predicate,
                    "confidence": confidence,
                    "context": context[:500],
                    "mention_count": 0,
                    "methods": set(),
                    "evidence": [],
                    "confidence_breakdown": confidence_breakdown,
                },
            )
            record["mention_count"] += 1
            if confidence > record["confidence"]:
                record["confidence"] = confidence
                record["confidence_breakdown"] = confidence_breakdown
            if not record["context"] and context:
                record["context"] = context[:500]
            if context and context not in record["evidence"] and len(record["evidence"]) < 5:
                record["evidence"].append(context[:500])
            record["methods"].add(method)

    entity_values = sorted(entity_records.values(), key=lambda item: (-item["confidence"], item["text"]))
    relation_values = sorted(relation_records.values(), key=lambda item: (-item["confidence"], item["predicate"]))
    if len(entity_values) > _MAX_ENTITIES:
        warnings.append(f"Entity candidates were capped at {_MAX_ENTITIES}.")
        entity_values = entity_values[:_MAX_ENTITIES]
    allowed_entity_ids = {item["id"] for item in entity_values}
    relation_values = [
        item for item in relation_values
        if item["source_id"] in allowed_entity_ids and item["target_id"] in allowed_entity_ids
    ]
    if len(relation_values) > _MAX_RELATIONS:
        warnings.append(f"Relation candidates were capped at {_MAX_RELATIONS}.")
        relation_values = relation_values[:_MAX_RELATIONS]

    entity_candidates = [
        DocumentEntityCandidate(
            id=item["id"],
            text=item["text"],
            label=item["label"],
            confidence=item["confidence"],
            mention_count=item["mention_count"],
            properties={
                "offsets": item["offsets"],
                "aliases": sorted(item["aliases"]),
                "evidence": item["evidence"],
                "extraction_methods": sorted(item["methods"]),
                "confidence_breakdown": item["confidence_breakdown"],
            },
        )
        for item in entity_values
    ]
    relation_candidates = [
        DocumentRelationCandidate(
            id=item["id"],
            source_id=item["source_id"],
            target_id=item["target_id"],
            predicate=item["predicate"],
            confidence=item["confidence"],
            context=item["context"],
            mention_count=item["mention_count"],
            properties={
                "extraction_methods": sorted(item["methods"]),
                "evidence": item["evidence"],
                "confidence_breakdown": item["confidence_breakdown"],
            },
        )
        for item in relation_values
    ]
    return entity_candidates, relation_candidates, sorted(methods_used), list(dict.fromkeys(warnings))


@router.post("/preview", response_model=DocumentPreviewResponse)
async def preview_document_graph(
    request: Request,
    file: UploadFile = File(...),
    language: str = Form("auto"),
    extraction_method: str = Form("auto"),
    min_confidence: float = Form(0.45),
    provider: str = Form("openai"),
    model: Optional[str] = Form(None),
    ontology_profile: str = Form("general"),
    custom_entity_types: Optional[str] = Form(None),
    custom_relation_types: Optional[str] = Form(None),
) -> DocumentPreviewResponse:
    filename = _safe_filename(file.filename)
    extension = Path(filename).suffix.lower()
    if extension not in _ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=422,
            detail=f"Unsupported document type '{extension}'. Allowed: {sorted(_ALLOWED_EXTENSIONS)}",
        )
    if language not in {"auto", "en", "zh"}:
        raise HTTPException(status_code=422, detail="language must be auto, en, or zh")
    if extraction_method not in {"auto", "rules", "llm"}:
        raise HTTPException(status_code=422, detail="extraction_method must be auto, rules, or llm")
    if not 0.0 <= min_confidence <= 1.0:
        raise HTTPException(status_code=422, detail="min_confidence must be between 0 and 1")
    provider = provider.strip().lower()
    if provider not in _SUPPORTED_PROVIDERS:
        raise HTTPException(status_code=422, detail=f"Unsupported LLM provider '{provider}'.")
    resolved_model = (model or "").strip() or _DEFAULT_MODELS[provider]
    resolved_profile, entity_types, relation_types = _resolve_ontology(
        ontology_profile,
        custom_entity_types,
        custom_relation_types,
    )

    content = await file.read(_DOCUMENT_MAX_BYTES + 1)
    if len(content) > _DOCUMENT_MAX_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"Document exceeds the {_DOCUMENT_MAX_BYTES // (1024 * 1024)} MB limit.",
        )
    if not content:
        raise HTTPException(status_code=422, detail="Document is empty.")

    text, parser_name, parser_warnings = await asyncio.to_thread(_parse_document, filename, content)
    if not text:
        raise HTTPException(status_code=422, detail="No extractable text was found in the document.")
    warnings = list(parser_warnings)
    if len(text) > _TEXT_MAX_CHARACTERS:
        text = text[:_TEXT_MAX_CHARACTERS]
        warnings.append(f"Document text was truncated to {_TEXT_MAX_CHARACTERS:,} characters.")

    resolved_language = _detect_language(text, language)
    chunks = _chunk_text(text)
    session_api_key = _session_credentials(request).get(provider)
    effective_api_key = session_api_key
    if not effective_api_key and provider != "ollama":
        from ...semantic_extract.config import config

        effective_api_key = config.get_api_key(provider)
    provider_status = await asyncio.to_thread(
        _provider_runtime_status,
        provider,
        resolved_model,
        session_api_key,
    )
    if extraction_method == "llm" and not provider_status.available:
        raise HTTPException(
            status_code=503,
            detail=(
                f"LLM provider '{provider}' is unavailable: {provider_status.reason} "
                "Configure the server-side credentials/dependency or select another mode."
            ),
        )
    if extraction_method == "auto" and not provider_status.available:
        warnings.append(
            f"{provider} is unavailable; Automatic mode selected an offline extractor."
        )
    provider_ready = provider_status.available
    provider_instance: Optional[Any] = None
    if extraction_method in {"auto", "llm"} and provider_ready:
        try:
            provider_instance = await asyncio.to_thread(
                _create_request_provider,
                provider,
                resolved_model,
                effective_api_key,
            )
        except Exception as exc:
            code, retryable, upstream_status = _classify_llm_error(exc)
            initialization_failure = LLMExtractionFailure(
                code="provider_initialization_error" if code == "provider_structured_output_error" else code,
                provider=provider,
                model=resolved_model,
                phase="provider_initialization",
                chunk_index=1,
                total_chunks=len(chunks),
                attempts=1,
                retryable=retryable,
                upstream_status=upstream_status,
            )
            if extraction_method == "llm":
                raise HTTPException(status_code=502, detail=initialization_failure.detail()) from exc
            provider_ready = False
            warnings.append(
                f"{provider} initialization failed; Automatic mode selected an offline extractor."
            )
    try:
        try:
            entities, relations, actual_methods, extraction_warnings = await asyncio.to_thread(
                _merge_candidates,
                chunks,
                resolved_language,
                extraction_method,
                min_confidence,
                provider,
                resolved_model,
                provider_ready,
                provider_instance,
                entity_types,
                relation_types,
            )
        except LLMExtractionFailure as exc:
            raise HTTPException(status_code=502, detail=exc.detail()) from exc
    finally:
        await asyncio.to_thread(_close_request_provider, provider_instance)
    warnings.extend(extraction_warnings)
    if not entities:
        warnings.append("No entity candidates were extracted. Try LLM mode or a lower confidence threshold.")
    elif not relations:
        warnings.append("Entities were found, but no relation candidates were extracted.")

    media_type = file.content_type or "application/octet-stream"
    document_id = hashlib.sha256(content).hexdigest()
    return DocumentPreviewResponse(
        document_id=document_id,
        filename=filename,
        media_type=media_type,
        parser=parser_name,
        language=resolved_language,
        extraction_method=", ".join(actual_methods) or extraction_method,
        execution=ExtractionExecution(
            requested_method=extraction_method,
            actual_methods=actual_methods,
            provider=provider if extraction_method in {"auto", "llm"} else None,
            model=resolved_model if extraction_method in {"auto", "llm"} else None,
            provider_available=provider_ready,
            fallback_used=extraction_method == "auto" and "llm-structured" not in actual_methods,
            ontology_profile=resolved_profile,
        ),
        ontology_profile=resolved_profile,
        character_count=len(text),
        chunk_count=len(chunks),
        text_preview=text[:3_000],
        entities=entities,
        relations=relations,
        warnings=warnings,
    )


@router.post("/commit", response_model=DocumentCommitResponse)
async def commit_document_graph(
    body: DocumentCommitRequest,
    session: GraphSession = Depends(get_session),
) -> DocumentCommitResponse:
    if not body.entities:
        raise HTTPException(status_code=422, detail="Select at least one entity before committing.")

    document_node_id = f"document:{body.document_id[:32]}"
    source_properties = {
        "source_document": body.filename,
        "source_document_id": body.document_id,
        "document_node_id": document_node_id,
        "parser": body.parser,
        "language": body.language,
        "extraction_method": body.extraction_method,
        "ontology_profile": body.ontology_profile,
    }
    nodes: List[Dict[str, Any]] = [
        {
            "id": document_node_id,
            "type": "DOCUMENT",
            "properties": {
                "content": body.filename,
                "filename": body.filename,
                "media_type": body.media_type,
                "character_count": body.character_count,
                **source_properties,
            },
        }
    ]
    selected_entity_ids = {entity.id for entity in body.entities}
    for entity in body.entities:
        nodes.append(
            {
                "id": entity.id,
                "type": entity.label or "ENTITY",
                "properties": {
                    "content": entity.text,
                    "confidence": entity.confidence,
                    "mention_count": entity.mention_count,
                    **entity.properties,
                    **source_properties,
                },
            }
        )

    edges: List[Dict[str, Any]] = []
    for relation in body.relations:
        if relation.source_id not in selected_entity_ids or relation.target_id not in selected_entity_ids:
            continue
        edges.append(
            {
                "id": relation.id,
                "source_id": relation.source_id,
                "target_id": relation.target_id,
                "type": relation.predicate or "related_to",
                "weight": relation.confidence,
                "properties": {
                    "confidence": relation.confidence,
                    "context": relation.context,
                    "mention_count": relation.mention_count,
                    **relation.properties,
                    **source_properties,
                },
            }
        )

    for entity in body.entities:
        provenance_key = f"{entity.id}\0extracted_from\0{document_node_id}".encode("utf-8")
        edges.append(
            {
                "id": f"provenance:{hashlib.sha1(provenance_key).hexdigest()[:16]}",
                "source_id": entity.id,
                "target_id": document_node_id,
                "type": "extracted_from",
                "weight": entity.confidence,
                "properties": {
                    "confidence": entity.confidence,
                    "mention_count": entity.mention_count,
                    **source_properties,
                },
            }
        )

    try:
        nodes_added, edges_added = await asyncio.to_thread(session.add_nodes_and_edges, nodes, edges)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    return DocumentCommitResponse(
        document_node_id=document_node_id,
        nodes_added=nodes_added,
        edges_added=edges_added,
        entities_submitted=len(body.entities),
        relations_submitted=sum(
            1 for relation in body.relations
            if relation.source_id in selected_entity_ids and relation.target_id in selected_entity_ids
        ),
    )
